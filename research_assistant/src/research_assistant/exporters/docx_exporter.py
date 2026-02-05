"""DOCX exporter for markdown-ish reports.

The report content produced by the agents is markdown. This exporter implements a
small subset (headings + bullets + numbered lists + paragraphs) to create a
presentable Word document without requiring Pandoc.
"""

from __future__ import annotations

from dataclasses import dataclass
import re
from typing import Iterable


@dataclass(frozen=True)
class DocxExportOptions:
    title: str
    subtitle: str | None = None


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _iter_lines(text: str) -> Iterable[str]:
    for line in (text or "").splitlines():
        yield line.rstrip("\n")


def _add_markdown_inline_runs(paragraph, text: str) -> None:
    """Add runs, making **bold** segments bold."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        bold_run = paragraph.add_run(m.group(1))
        bold_run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def export_docx(markdown_text: str, output_path: str, *, options: DocxExportOptions) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency for DOCX export. Install with:\n"
            "  pip install python-docx\n"
            "Or:\n"
            "  pip install -e '.[export]'"
        ) from e

    doc = Document()

    # Title block
    doc.add_heading(options.title, level=0)
    if options.subtitle:
        doc.add_paragraph(options.subtitle)

    doc.add_paragraph("")  # spacer

    for raw in _iter_lines(markdown_text):
        line = raw.strip()
        if not line:
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        # Bullets
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_markdown_inline_runs(p, line[2:].strip())
            continue

        # Numbered list (simple)
        m = re.match(r"^(\d+)[\.\)]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_markdown_inline_runs(p, m.group(2).strip())
            continue

        # Paragraph
        p = doc.add_paragraph()
        _add_markdown_inline_runs(p, line)

    doc.save(output_path)

